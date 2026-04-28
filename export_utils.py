from __future__ import annotations

import io
import json
from typing import Dict, List, Any

import pandas as pd

from checklist_pipeline import SECTIONS


def checklist_to_df(checklist: Dict[str, List[Any]]) -> pd.DataFrame:
    rows = []
    for section in SECTIONS:
        for tc in checklist.get(section, []):
            if isinstance(tc, dict):
                rows.append({
                    "Category": section,
                    "Id": tc.get("id", ""),
                    "Test Summary": tc.get("test_summary", ""),
                    "Description": tc.get("description", ""),
                    "Action": tc.get("action", ""),
                    "Data": tc.get("data", "N/A"),
                    "Expected Result": tc.get("expected_result", ""),
                })
            else:
                rows.append({
                    "Category": section,
                    "Id": "",
                    "Test Summary": str(tc),
                    "Description": "",
                    "Action": "",
                    "Data": "N/A",
                    "Expected Result": "",
                })
    return pd.DataFrame(rows)


def checklist_to_excel_bytes(checklist: Dict[str, List[Any]]) -> bytes:
    df = checklist_to_df(checklist)
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Test Cases")
        ws = writer.sheets["Test Cases"]
        for col in ws.columns:
            max_len = max(len(str(cell.value or "")) for cell in col)
            ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 60)
    return buf.getvalue()


def checklist_to_docx_bytes(checklist: Dict[str, List[Any]], title: str = "Test Checklist") -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return b""

    doc = Document()
    doc.add_heading(title, 0)

    for section in SECTIONS:
        items = checklist.get(section, [])
        if not items:
            continue
        doc.add_heading(section, level=1)
        for tc in items:
            if isinstance(tc, dict):
                doc.add_heading(f"{tc.get('id','')} — {tc.get('test_summary','')}", level=2)
                table = doc.add_table(rows=5, cols=2)
                table.style = "Table Grid"
                fields = [
                    ("Description", tc.get("description", "")),
                    ("Action", tc.get("action", "")),
                    ("Data", tc.get("data", "N/A")),
                    ("Expected Result", tc.get("expected_result", "")),
                ]
                for i, (label, value) in enumerate(fields):
                    table.rows[i].cells[0].text = label
                    table.rows[i].cells[1].text = value
                table.rows[4].cells[0].text = "Category"
                table.rows[4].cells[1].text = section
                doc.add_paragraph("")
            else:
                doc.add_paragraph(str(tc), style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def checklist_to_jira_json_bytes(checklist: Dict[str, List[Any]], summary: str = "Test Checklist") -> bytes:
    issues = []
    for section in SECTIONS:
        for tc in checklist.get(section, []):
            if isinstance(tc, dict):
                issues.append({
                    "fields": {
                        "summary": tc.get("test_summary", ""),
                        "description": (
                            f"*Description:* {tc.get('description','')}\n\n"
                            f"*Action:* {tc.get('action','')}\n\n"
                            f"*Data:* {tc.get('data','N/A')}\n\n"
                            f"*Expected Result:* {tc.get('expected_result','')}"
                        ),
                        "issuetype": {"name": "Test"},
                        "labels": [section, "AI-Generated"],
                        "customfield_test_id": tc.get("id", ""),
                    }
                })
            else:
                issues.append({
                    "fields": {
                        "summary": str(tc),
                        "issuetype": {"name": "Test"},
                        "labels": [section, "AI-Generated"],
                    }
                })
    payload = {"issueUpdates": issues}
    return json.dumps(payload, indent=2, ensure_ascii=False).encode("utf-8")


def checklist_to_testrail_csv_bytes(checklist: Dict[str, List[Any]]) -> bytes:
    rows = []
    for section in SECTIONS:
        for tc in checklist.get(section, []):
            if isinstance(tc, dict):
                rows.append({
                    "ID": tc.get("id", ""),
                    "Title": tc.get("test_summary", ""),
                    "Section": section,
                    "Type": "Functional",
                    "Priority": "Medium",
                    "Description": tc.get("description", ""),
                    "Action": tc.get("action", ""),
                    "Data": tc.get("data", "N/A"),
                    "Expected Result": tc.get("expected_result", ""),
                })
            else:
                rows.append({
                    "ID": "",
                    "Title": str(tc),
                    "Section": section,
                    "Type": "Functional",
                    "Priority": "Medium",
                    "Description": "",
                    "Action": "",
                    "Data": "N/A",
                    "Expected Result": "",
                })
    df = pd.DataFrame(rows)
    return df.to_csv(index=False).encode("utf-8")